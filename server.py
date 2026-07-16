"""Resume-Tailoring MCP Server.

An MCP server that helps an AI client (Claude Desktop, Cursor, etc.) tailor a
person's CV to a specific job. The MCP does NOT rewrite the CV — Claude does
that. The MCP provides the parts a plain chat can't do well:

  * Persist a master CV between sessions.
  * Pull a job posting from a URL (or pasted text).
  * Run a deterministic ATS keyword check (real math, not vibes).
  * Export the finished CV to a clean, ATS-safe PDF or DOCX.

Flow:  load_master_resume + fetch_job_posting + extract_keywords
       -> (Claude rewrites the CV) -> ats_gap_check -> export_resume
"""

from __future__ import annotations

import json
import os
import re
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Any, Literal

import httpx
from bs4 import BeautifulSoup
from mcp.server.fastmcp import FastMCP

mcp = FastMCP("resume-tailor")


# --------------------------------------------------------------------------- #
# Shared helpers — one place for file I/O and one for HTTP.
# --------------------------------------------------------------------------- #

def _store_path() -> Path:
    """Resolve the master-CV JSON path (RESUME_STORE_PATH overrides default)."""
    raw = os.environ.get("RESUME_STORE_PATH", "~/.resume-mcp/master.json")
    return Path(raw).expanduser()


def _export_dir() -> Path:
    """Directory where exported resumes are written (next to the store)."""
    d = _store_path().parent / "exports"
    d.mkdir(parents=True, exist_ok=True)
    return d


def read_json(path: Path) -> dict[str, Any]:
    """Read a JSON file, raising a clear error if it is missing or malformed."""
    if not path.exists():
        raise FileNotFoundError(
            f"No master resume found at {path}. "
            f"Call save_master_resume first to store your base CV."
        )
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as exc:
        raise ValueError(f"Master resume at {path} is not valid JSON: {exc}") from exc


def write_json(path: Path, data: dict[str, Any]) -> None:
    """Write JSON to disk, creating parent directories as needed."""
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")


def fetch_url(url: str) -> str:
    """Fetch a URL and return raw HTML, with a clear, readable error on failure."""
    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/122.0 Safari/537.36"
        )
    }
    with httpx.Client(follow_redirects=True, timeout=20.0, headers=headers) as client:
        resp = client.get(url)
        resp.raise_for_status()
        return resp.text


# --------------------------------------------------------------------------- #
# Keyword extraction data — deterministic, no ML.
# --------------------------------------------------------------------------- #

# Filler words we never want to surface as "skills".
STOPLIST: set[str] = {
    "a", "an", "the", "and", "or", "but", "if", "then", "else", "for", "to",
    "of", "in", "on", "at", "by", "with", "as", "is", "are", "was", "were",
    "be", "been", "being", "this", "that", "these", "those", "it", "its",
    "you", "your", "we", "our", "they", "their", "he", "she", "his", "her",
    "will", "would", "should", "can", "could", "may", "might", "must", "have",
    "has", "had", "do", "does", "did", "not", "no", "yes", "from", "into",
    "about", "who", "what", "when", "where", "why", "how", "all", "any", "some",
    "more", "most", "other", "such", "only", "own", "same", "so", "than", "too",
    "very", "just", "up", "out", "off", "over", "under", "again", "role",
    "team", "work", "working", "years", "year", "experience", "job", "company",
    "candidate", "position", "opportunity", "join", "help", "looking", "strong",
    "excellent", "good", "great", "ability", "skills", "skill", "including",
    "etc", "e", "g", "ie", "us", "well", "new", "using", "use", "used", "across",
    "within", "per", "via", "plus", "also", "both", "each", "many", "much",
    "world", "people", "make", "made", "build", "get", "like", "want", "need",
    "day", "days", "time", "part", "full", "based", "benefits", "salary",
    "apply", "please", "email", "contact", "location", "remote", "hybrid",
    "onsite", "office", "requirements", "responsibilities", "qualifications",
    "preferred", "required", "nice", "must-have", "you'll", "we're", "we'll",
    "familiarity", "familiar", "collaborate", "collaborating", "understanding",
    "knowledge", "proficiency", "proficient", "senior", "junior", "mid",
    "responsible", "ensure", "deliver", "develop", "developing", "designers",
}

# Known multi-word skills/tools -> canonical display casing. Matched as phrases
# before the single-token pass so "REST APIs" isn't split into "rest" + "apis".
# Several source spellings map to one display so the final list de-duplicates
# (e.g. "rest api" and "rest apis" both collapse to "REST APIs").
KNOWN_PHRASES: dict[str, str] = {
    "rest apis": "REST APIs", "rest api": "REST APIs", "restful apis": "REST APIs",
    "graphql": "GraphQL", "machine learning": "Machine Learning",
    "deep learning": "Deep Learning",
    "natural language processing": "Natural Language Processing",
    "computer vision": "Computer Vision", "data science": "Data Science",
    "data analysis": "Data Analysis", "data engineering": "Data Engineering",
    "unit testing": "Unit Testing",
    "test driven development": "Test-Driven Development",
    "continuous integration": "CI/CD", "continuous deployment": "CI/CD",
    "ci/cd": "CI/CD", "version control": "Version Control",
    "object oriented": "Object-Oriented", "design patterns": "Design Patterns",
    "microservices": "Microservices", "message queues": "Message Queues",
    "event driven": "Event-Driven", "distributed systems": "Distributed Systems",
    "cloud computing": "Cloud Computing",
    "infrastructure as code": "Infrastructure as Code",
    "site reliability": "Site Reliability", "agile": "Agile", "scrum": "Scrum",
    "project management": "Project Management",
    "product management": "Product Management", "google cloud": "Google Cloud",
    "amazon web services": "AWS", "node.js": "Node.js", "next.js": "Next.js",
    "vue.js": "Vue.js", "react native": "React Native", "spring boot": "Spring Boot",
    "ruby on rails": "Ruby on Rails", "asp.net": "ASP.NET",
    "github actions": "GitHub Actions", "gitlab ci": "GitLab CI",
    "material ui": "Material UI", "tailwind css": "Tailwind CSS",
    "user experience": "User Experience", "user interface": "User Interface",
    "front end": "Front-End", "back end": "Back-End", "full stack": "Full-Stack",
    "responsive design": "Responsive Design", "web accessibility": "Web Accessibility",
    "a/b testing": "A/B Testing",
}

# Known single-token skills/tools — canonical display casing.
KNOWN_TERMS: dict[str, str] = {
    t.lower(): t
    for t in [
        "Python", "JavaScript", "TypeScript", "Java", "Go", "Rust", "C", "C++",
        "C#", "Ruby", "PHP", "Swift", "Kotlin", "Scala", "R", "MATLAB", "SQL",
        "HTML", "CSS", "SASS", "React", "Angular", "Vue", "Svelte", "Redux",
        "jQuery", "Bootstrap", "Django", "Flask", "FastAPI", "Express",
        "Node", "Deno", "Rails", "Laravel", "Spring", "Kafka", "RabbitMQ",
        "Redis", "PostgreSQL", "Postgres", "MySQL", "MongoDB", "SQLite",
        "DynamoDB", "Cassandra", "Elasticsearch", "Snowflake", "Databricks",
        "Docker", "Kubernetes", "Terraform", "Ansible", "Jenkins", "AWS",
        "Azure", "GCP", "Heroku", "Vercel", "Netlify", "Linux", "Bash",
        "Git", "GitHub", "GitLab", "Bitbucket", "Jira", "Confluence",
        "Figma", "Sketch", "Photoshop", "Illustrator", "Pandas", "NumPy",
        "PyTorch", "TensorFlow", "Keras", "Scikit-learn", "Spark", "Hadoop",
        "Airflow", "Tableau", "PowerBI", "Excel", "GraphQL", "REST", "gRPC",
        "OAuth", "JWT", "Nginx", "Apache", "Webpack", "Vite", "Babel",
        "Jest", "Pytest", "Selenium", "Cypress", "Playwright", "Prometheus",
        "Grafana", "Datadog", "Sentry", "Stripe", "Twilio", "GraphQL",
        "Salesforce", "SAP", "Kanban", "Agile", "Scrum", "DevOps", "SRE",
        "SEO", "SDK", "ETL", "ORM", "SaaS",
        "NoSQL", "OOP", "TDD", "UX", "UI", "SOLID", "Numpy", "Matplotlib",
    ]
}

# Section headers whose contents weigh more heavily (ATS scans these hardest).
HEAVY_SECTION_RE = re.compile(
    r"(requirements|qualifications|skills|technologies|tech stack|"
    r"what you.?ll need|must have|responsibilities|experience with|"
    r"we.?re looking for)",
    re.IGNORECASE,
)

_TOKEN_RE = re.compile(r"[A-Za-z][A-Za-z0-9+#.\-]*")


def _normalize(text: str) -> str:
    return re.sub(r"\s+", " ", text.lower()).strip()


def rank_keywords(job_text: str, top_n: int = 30) -> list[dict[str, Any]]:
    """Deterministically pull ranked skill/tool keywords from job text.

    Strategy: score by frequency, boost terms appearing in heavy sections
    (skills/requirements) and known-skill terms. Returns ranked descriptors.
    """
    norm = _normalize(job_text)

    # Identify which characters fall inside a "heavy" section for weighting.
    heavy_spans: list[tuple[int, int]] = []
    lines = job_text.splitlines()
    offset = 0
    heavy = False
    for line in lines:
        start = offset
        offset += len(line) + 1
        stripped = line.strip()
        if HEAVY_SECTION_RE.search(stripped) and len(stripped) < 60:
            heavy = True
            continue
        # A blank line or a new short header ends a heavy block.
        if heavy and not stripped:
            heavy = False
        if heavy:
            heavy_spans.append((start, offset))

    def in_heavy(idx: int) -> bool:
        return any(s <= idx < e for s, e in heavy_spans)

    scores: Counter[str] = Counter()
    display: dict[str, str] = {}

    # 1) Multi-word known phrases (several spellings collapse to one display).
    for phrase, shown in KNOWN_PHRASES.items():
        for m in re.finditer(re.escape(phrase), norm):
            weight = 3 + (2 if in_heavy(m.start()) else 0)
            scores[shown.lower()] += weight
            display.setdefault(shown.lower(), shown)

    # 2) Single tokens.
    for m in _TOKEN_RE.finditer(job_text):
        raw = m.group(0)
        low = raw.lower().strip(".-")
        if not low or low in STOPLIST or len(low) < 2:
            continue
        known = low in KNOWN_TERMS
        # Skip generic lowercase words unless they're known skills.
        if not known and (low.isalpha() and raw[0].islower() and len(low) > 2):
            continue
        weight = 1
        if known:
            weight += 2
        if in_heavy(m.start()):
            weight += 2
        scores[low] += weight
        if known:
            display[low] = KNOWN_TERMS[low]
        else:
            display.setdefault(low, raw)

    known_display = {v.lower() for v in KNOWN_PHRASES.values()} | set(KNOWN_TERMS)
    ranked = scores.most_common()
    out: list[dict[str, Any]] = []
    seen_display: set[str] = set()
    covered_words: set[str] = set()  # words already inside an accepted phrase
    _word_re = re.compile(r"[a-z0-9]+")
    for term, score in ranked:
        name = display.get(term, term)
        key = name.lower()
        if key in seen_display:
            continue
        words = _word_re.findall(key)
        # Suppress a single-word token already covered by an accepted phrase
        # (e.g. drop "REST"/"CI" once "REST APIs"/"CI/CD" are in).
        if len(words) == 1 and words[0] in covered_words:
            continue
        is_known = key in known_display or term in KNOWN_TERMS
        # Drop one-off unknown capitalized words: keep an unknown term only if
        # it recurs or appears in a skills/requirements section (score >= 2).
        if not is_known and score < 2:
            continue
        seen_display.add(key)
        if len(words) > 1:
            covered_words.update(words)
        out.append({"keyword": name, "score": score, "known_skill": is_known})
        if len(out) >= top_n:
            break
    return out


# --------------------------------------------------------------------------- #
# Tool 1 & 2 — persist and load the master CV.
# --------------------------------------------------------------------------- #

@mcp.tool()
def save_master_resume(resume: dict[str, Any]) -> str:
    """Store or update the user's base CV so it persists between sessions.

    Provide the CV as structured JSON. Recommended keys:
      contact    - {name, email, phone, location, links}
      summary    - short professional summary (string)
      experience - list of {title, company, location, start, end, bullets[]}
      projects   - list of {name, description, tech[], link}
      skills     - list of strings (or {category, items[]})
      education  - list of {degree, school, location, start, end, details}

    This is step 1 of the flow: set it up once, reuse it for every job.
    Returns a confirmation with the stored path.
    """
    path = _store_path()
    write_json(path, resume)
    name = resume.get("contact", {}).get("name", "resume")
    return f"Saved master resume for '{name}' to {path}."


@mcp.tool()
def load_master_resume() -> dict[str, Any]:
    """Return the stored master CV so Claude can tailor it to a job.

    This is the starting point of every tailoring session. If no resume is
    stored yet, raises a clear error telling the user to save one first.
    """
    return read_json(_store_path())


# --------------------------------------------------------------------------- #
# Tool 3 — fetch a job posting (URL -> clean text, with paste fallback).
# --------------------------------------------------------------------------- #

@mcp.tool()
def fetch_job_posting(url: str = "", pasted_text: str = "") -> dict[str, Any]:
    """Get a job posting as clean readable text.

    Pass a `url` to fetch and strip a posting to plain text, OR pass
    `pasted_text` directly if you already have the description. URL fetching is
    best-effort: if the site is blocked or login-walled, this returns a clear
    message asking the user to paste the text instead — it never fails silently.

    Returns {source, text, char_count}.
    """
    if pasted_text.strip():
        text = re.sub(r"\n{3,}", "\n\n", pasted_text.strip())
        return {"source": "pasted", "text": text, "char_count": len(text)}

    if not url.strip():
        return {
            "source": "none",
            "text": "",
            "char_count": 0,
            "message": "Provide a job `url` to fetch, or paste the description via `pasted_text`.",
        }

    try:
        html = fetch_url(url)
    except Exception as exc:  # noqa: BLE001 - surface a clean message to Claude
        return {
            "source": "url_failed",
            "text": "",
            "char_count": 0,
            "message": (
                f"Could not fetch {url} ({type(exc).__name__}: {exc}). "
                f"The site may be blocked or login-walled — please paste the "
                f"job description text directly via `pasted_text` instead."
            ),
        }

    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header", "form", "noscript"]):
        tag.decompose()
    text = soup.get_text(separator="\n")
    lines = [ln.strip() for ln in text.splitlines()]
    text = "\n".join(ln for ln in lines if ln)
    text = re.sub(r"\n{3,}", "\n\n", text)

    if len(text) < 200:
        return {
            "source": "url_thin",
            "text": text,
            "char_count": len(text),
            "message": (
                f"Fetched {url} but extracted very little text ({len(text)} chars) — "
                f"the posting is likely rendered by JavaScript or login-walled. "
                f"Please paste the job description via `pasted_text` instead."
            ),
        }

    return {"source": "url", "url": url, "text": text, "char_count": len(text)}


# --------------------------------------------------------------------------- #
# Tool 4 — deterministic keyword extraction.
# --------------------------------------------------------------------------- #

@mcp.tool()
def extract_keywords(job_text: str, top_n: int = 30) -> dict[str, Any]:
    """Pull the skills/tools/keywords an ATS would scan for from a job posting.

    Deterministic (no ML): tokenizes the text, filters filler words, recognizes
    known skills and multi-word phrases (e.g. "REST APIs"), and ranks by
    frequency — boosting terms that appear in skills/requirements sections.

    Returns {keywords: [ranked strings], detail: [{keyword, score, known_skill}]}.
    Use the ranked keyword list as input to ats_gap_check.
    """
    detail = rank_keywords(job_text, top_n=top_n)
    return {
        "keywords": [d["keyword"] for d in detail],
        "detail": detail,
        "count": len(detail),
    }


# --------------------------------------------------------------------------- #
# Tool 5 — the ATS gap check (the standout feature).
# --------------------------------------------------------------------------- #

def _term_present(term: str, haystack: str) -> bool:
    """Whole-word-ish presence check, tolerant of plurals and separators."""
    t = term.lower().strip()
    if not t:
        return False
    # Escape, then allow flexible whitespace/punctuation between phrase words.
    parts = [re.escape(p) for p in re.split(r"\s+", t)]
    pattern = r"[\s\-/]*".join(parts)
    # Word-ish boundaries; allow trailing 's' for simple plurals.
    regex = re.compile(rf"(?<![A-Za-z0-9]){pattern}s?(?![A-Za-z])", re.IGNORECASE)
    return bool(regex.search(haystack))


@mcp.tool()
def ats_gap_check(resume_text: str, keywords: list[str]) -> dict[str, Any]:
    """Compare a CV against job keywords → match score (%) + missing terms.

    This is the tool's killer feature: it tells the user *concretely* what to
    add. Pass the resume as plain text and the ranked keyword list from
    extract_keywords. Returns which keywords are present, which are missing, and
    a match percentage — so Claude knows exactly what to surface in the rewrite.

    Returns {match_score, matched, missing, total_keywords}.
    """
    haystack = resume_text.lower()
    matched: list[str] = []
    missing: list[str] = []
    for kw in keywords:
        if _term_present(kw, haystack):
            matched.append(kw)
        else:
            missing.append(kw)

    total = len(keywords)
    score = round(100 * len(matched) / total, 1) if total else 0.0
    return {
        "match_score": score,
        "matched": matched,
        "missing": missing,
        "total_keywords": total,
        "summary": (
            f"{len(matched)}/{total} keywords matched ({score}%). "
            + (
                f"Missing: {', '.join(missing)}."
                if missing
                else "No gaps — strong keyword coverage."
            )
        ),
    }


# --------------------------------------------------------------------------- #
# Tool 6 — export to a clean, ATS-safe PDF or DOCX.
# --------------------------------------------------------------------------- #

def _blocks_from_content(content: str | dict[str, Any]) -> list[tuple[str, str]]:
    """Turn markdown text OR structured CV JSON into simple (type, text) blocks.

    Block types: 'h1', 'h2', 'bullet', 'para'. Deliberately flat and
    single-column so ATS parsers can read every line as real text.
    """
    blocks: list[tuple[str, str]] = []

    if isinstance(content, dict):
        contact = content.get("contact", {})
        if contact.get("name"):
            blocks.append(("h1", contact["name"]))
        line = " | ".join(
            str(contact[k])
            for k in ("email", "phone", "location")
            if contact.get(k)
        )
        links = contact.get("links") or {}
        if isinstance(links, dict):
            line = " | ".join(filter(None, [line, *[str(v) for v in links.values()]]))
        if line:
            blocks.append(("para", line))
        if content.get("summary"):
            blocks.append(("h2", "Summary"))
            blocks.append(("para", content["summary"]))
        if content.get("experience"):
            blocks.append(("h2", "Experience"))
            for job in content["experience"]:
                head = " — ".join(filter(None, [job.get("title"), job.get("company")]))
                dates = " ".join(filter(None, [job.get("start"), job.get("end")])).strip()
                if dates:
                    head = f"{head} ({dates})"
                blocks.append(("h3", head))
                for b in job.get("bullets", []):
                    blocks.append(("bullet", b))
        if content.get("projects"):
            blocks.append(("h2", "Projects"))
            for p in content["projects"]:
                head = p.get("name", "")
                tech = ", ".join(p.get("tech", []))
                if tech:
                    head = f"{head} — {tech}"
                blocks.append(("h3", head))
                if p.get("description"):
                    blocks.append(("bullet", p["description"]))
        if content.get("skills"):
            blocks.append(("h2", "Skills"))
            skills = content["skills"]
            if skills and isinstance(skills[0], dict):
                for grp in skills:
                    items = ", ".join(grp.get("items", []))
                    blocks.append(("para", f"{grp.get('category', '')}: {items}"))
            else:
                blocks.append(("para", ", ".join(map(str, skills))))
        if content.get("education"):
            blocks.append(("h2", "Education"))
            for ed in content["education"]:
                head = " — ".join(filter(None, [ed.get("degree"), ed.get("school")]))
                dates = " ".join(filter(None, [ed.get("start"), ed.get("end")])).strip()
                if dates:
                    head = f"{head} ({dates})"
                blocks.append(("h3", head))
                if ed.get("details"):
                    blocks.append(("bullet", ed["details"]))
        return blocks

    # Markdown-ish text.
    for raw in content.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("# "):
            blocks.append(("h1", line[2:].strip()))
        elif line.startswith("## "):
            blocks.append(("h2", line[3:].strip()))
        elif line.startswith("### "):
            blocks.append(("h3", line[4:].strip()))
        elif line.lstrip().startswith(("- ", "* ")):
            blocks.append(("bullet", line.lstrip()[2:].strip()))
        else:
            blocks.append(("para", line.strip()))
    return blocks


def _safe_name(name: str) -> str:
    return re.sub(r"[^A-Za-z0-9_.-]+", "_", name).strip("_") or "resume"


def _export_docx(blocks: list[tuple[str, str]], out: Path) -> None:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for kind, text in blocks:
        if kind == "h1":
            p = doc.add_heading(text, level=0)
        elif kind == "h2":
            p = doc.add_heading(text, level=1)
        elif kind == "h3":
            p = doc.add_heading(text, level=2)
        elif kind == "bullet":
            p = doc.add_paragraph(text, style="List Bullet")
        else:
            p = doc.add_paragraph(text)
    doc.save(out)


def _export_pdf(blocks: list[tuple[str, str]], out: Path) -> None:
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        ListFlowable,
        ListItem,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
    )

    styles = getSampleStyleSheet()
    base = ParagraphStyle(
        "Body", parent=styles["Normal"], fontName="Helvetica",
        fontSize=10.5, leading=14, alignment=TA_LEFT,
    )
    h1 = ParagraphStyle("H1", parent=base, fontName="Helvetica-Bold", fontSize=18, leading=22, spaceAfter=2)
    h2 = ParagraphStyle("H2", parent=base, fontName="Helvetica-Bold", fontSize=13, leading=16, spaceBefore=10, spaceAfter=4)
    h3 = ParagraphStyle("H3", parent=base, fontName="Helvetica-Bold", fontSize=11, leading=14, spaceBefore=6, spaceAfter=1)

    doc = SimpleDocTemplate(
        str(out), pagesize=letter,
        leftMargin=0.7 * inch, rightMargin=0.7 * inch,
        topMargin=0.7 * inch, bottomMargin=0.7 * inch,
    )
    flow: list[Any] = []
    bullets: list[ListItem] = []

    def flush_bullets() -> None:
        if bullets:
            flow.append(ListFlowable(list(bullets), bulletType="bullet", leftIndent=14))
            bullets.clear()

    for kind, text in blocks:
        if kind == "bullet":
            bullets.append(ListItem(Paragraph(text, base), leftIndent=8))
            continue
        flush_bullets()
        if kind == "h1":
            flow.append(Paragraph(text, h1))
        elif kind == "h2":
            flow.append(Paragraph(text, h2))
        elif kind == "h3":
            flow.append(Paragraph(text, h3))
        else:
            flow.append(Paragraph(text, base))
    flush_bullets()
    doc.build(flow)


@mcp.tool()
def export_resume(
    content: str | dict[str, Any],
    format: Literal["pdf", "docx"] = "docx",
    filename: str = "",
) -> dict[str, Any]:
    """Render finished CV content to a clean, ATS-safe PDF or DOCX file.

    Pass `content` as either markdown text (use #/##/### headings and `- `
    bullets) OR structured CV JSON (same shape as the master resume). Choose
    `format` = "pdf" or "docx". The layout is deliberately single-column with
    standard fonts and real text (never image-rendered) so ATS parsers read it.

    Returns {path, format, blocks} with the saved file path.
    """
    blocks = _blocks_from_content(content)
    if not blocks:
        raise ValueError("No content to export — pass markdown text or CV JSON.")

    # Derive a filename from the first h1 (the name) if none given.
    if not filename:
        name = next((t for k, t in blocks if k == "h1"), "resume")
        stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
        filename = f"{_safe_name(name)}_{stamp}"
    filename = _safe_name(filename)

    out = _export_dir() / f"{filename}.{format}"
    if format == "docx":
        _export_docx(blocks, out)
    elif format == "pdf":
        _export_pdf(blocks, out)
    else:  # pragma: no cover - guarded by Literal
        raise ValueError(f"Unsupported format '{format}'. Use 'pdf' or 'docx'.")

    return {"path": str(out), "format": format, "blocks": len(blocks)}


@mcp.tool()
def export_cover_letter(
    content: str,
    format: Literal["pdf", "docx"] = "docx",
    filename: str = "",
    include_header: bool = True,
) -> dict[str, Any]:
    """Render a finished cover letter to a clean, ATS-safe PDF or DOCX file.

    The MCP does not write the letter — the client (Claude) writes it, tailored
    to the job using the master résumé and the ats_gap_check results. This tool
    just formats and saves it in the same single-column, real-text layout as the
    résumé so the two documents match.

    Pass `content` as the letter text (plain paragraphs, or markdown with
    `#`/`##` headings and `- ` bullets). When `include_header` is true and a
    master résumé is stored, the applicant's name + contact line are added at
    the top so the letter's letterhead matches the CV.

    Returns {path, format, blocks} with the saved file path.
    """
    blocks: list[tuple[str, str]] = []

    if include_header:
        try:
            master = read_json(_store_path())
        except (FileNotFoundError, ValueError):
            master = {}
        contact = master.get("contact", {}) if isinstance(master, dict) else {}
        if contact.get("name"):
            blocks.append(("h1", contact["name"]))
        line = " | ".join(
            str(contact[k])
            for k in ("email", "phone", "location")
            if contact.get(k)
        )
        if line:
            blocks.append(("para", line))

    blocks.extend(_blocks_from_content(content))
    if not blocks:
        raise ValueError("No content to export — pass the cover-letter text.")

    if not filename:
        name = next((t for k, t in blocks if k == "h1"), "cover_letter")
        stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
        filename = f"{_safe_name(name)}_cover_letter_{stamp}"
    filename = _safe_name(filename)

    out = _export_dir() / f"{filename}.{format}"
    if format == "docx":
        _export_docx(blocks, out)
    elif format == "pdf":
        _export_pdf(blocks, out)
    else:  # pragma: no cover - guarded by Literal
        raise ValueError(f"Unsupported format '{format}'. Use 'pdf' or 'docx'.")

    return {"path": str(out), "format": format, "blocks": len(blocks)}


if __name__ == "__main__":
    mcp.run()
